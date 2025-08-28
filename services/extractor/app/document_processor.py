import os
import io
import logging
import pandas as pd
import PyPDF2
from typing import List, Dict, Any, Optional
import aiofiles
import tempfile
import docx
import re
import pdfplumber
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process PDF and Word documents to extract text and data"""
    
    def __init__(self):
        self.supported_extensions = ('.pdf', '.doc', '.docx', '.csv', '.xls', '.xlsx')
        
    async def extract_text(self, content: bytes, filename: str) -> str:
        """Extract text from PDF or Word documents"""
        logger.info(f"Extracting text from {filename}")
        
        file_ext = os.path.splitext(filename.lower())[1]
        
        if file_ext == '.pdf':
            return await self._extract_from_pdf(content)
        elif file_ext in ('.doc', '.docx'):
            return await self._extract_from_word(content)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    async def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF using multiple methods for reliability"""
        text = ""
        
        # Try PyMuPDF first (usually best quality)
        try:
            with fitz.open(stream=content, filetype="pdf") as doc:
                text = ""
                for page in doc:
                    text += page.get_text()
                    
                if text.strip():
                    return text
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {str(e)}. Trying alternative method.")
        
        # Try pdfplumber next
        try:
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                temp.write(content)
                temp_path = temp.name
            
            with pdfplumber.open(temp_path) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
                text = "\n".join(pages)
                
            # Clean up temp file
            os.unlink(temp_path)
            
            if text.strip():
                return text
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}. Trying fallback method.")
        
        # Fallback to PyPDF2
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception as e:
            logger.error(f"All PDF extraction methods failed: {str(e)}")
            raise ValueError("Failed to extract text from PDF")
    
    async def _extract_from_word(self, content: bytes) -> str:
        """Extract text from Word document"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp:
                temp.write(content)
                temp_path = temp.name
            
            doc = docx.Document(temp_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += "\n" + row_text
            
            # Clean up temp file
            os.unlink(temp_path)
            
            return text
        except Exception as e:
            logger.error(f"Word extraction failed: {str(e)}")
            raise ValueError("Failed to extract text from Word document")
    
    async def process_csv(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process CSV, XLS, or XLSX file"""
        try:
            file_ext = os.path.splitext(filename.lower())[1]
            
            # Create temp file to use with pandas
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp:
                temp.write(content)
                temp_path = temp.name
            
            # Read with pandas based on file type
            if file_ext == '.csv':
                df = pd.read_csv(temp_path)
            elif file_ext in ('.xls', '.xlsx'):
                df = pd.read_excel(temp_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Clean up temp file
            os.unlink(temp_path)
            
            # Extract events
            events = self._extract_events_from_csv(df)
            
            # Extract vessel info
            vessel_info = self._extract_vessel_info_from_csv(df)
            
            return {
                "filename": filename,
                "vessel": vessel_info.get('vessel', 'Unknown'),
                "port": vessel_info.get('port', 'Unknown'),
                "cargo": vessel_info.get('cargo', 'Unknown'),
                "operation": vessel_info.get('operation', 'Discharge'),
                "voyage_from": vessel_info.get('voyage_from', 'Unknown'),
                "voyage_to": vessel_info.get('voyage_to', 'Unknown'),
                "events": events,
                "extracted_at": pd.Timestamp.now().isoformat(),
                "confidence_score": 0.9 if events else 0.5
            }
            
        except Exception as e:
            logger.error(f"Error processing {filename}: {str(e)}")
            raise ValueError(f"Failed to process {filename}: {str(e)}")
    
    def _extract_vessel_info_from_csv(self, df: pd.DataFrame) -> Dict[str, str]:
        """Extract vessel information from CSV"""
        vessel_info = {
            'vessel': 'Unknown',
            'port': 'Unknown',
            'cargo': 'Unknown',
            'operation': 'Discharge',
            'voyage_from': 'Unknown',
            'voyage_to': 'Unknown'
        }
        
        # Look for vessel info in column headers
        for col in df.columns:
            col_lower = col.lower()
            
            if 'vessel' in col_lower or 'ship' in col_lower:
                first_val = df[col].iloc[0] if not df.empty else None
                if pd.notna(first_val):
                    vessel_info['vessel'] = str(first_val)
            
            elif 'port' in col_lower:
                first_val = df[col].iloc[0] if not df.empty else None
                if pd.notna(first_val):
                    vessel_info['port'] = str(first_val)
            
            elif 'cargo' in col_lower:
                first_val = df[col].iloc[0] if not df.empty else None
                if pd.notna(first_val):
                    vessel_info['cargo'] = str(first_val)
            
            elif 'operation' in col_lower:
                first_val = df[col].iloc[0] if not df.empty else None
                if pd.notna(first_val):
                    vessel_info['operation'] = str(first_val)
            
            elif 'from' in col_lower:
                first_val = df[col].iloc[0] if not df.empty else None
                if pd.notna(first_val):
                    vessel_info['voyage_from'] = str(first_val)
            
            elif 'to' in col_lower:
                first_val = df[col].iloc[0] if not df.empty else None
                if pd.notna(first_val):
                    vessel_info['voyage_to'] = str(first_val)
        
        return vessel_info
    
    def _extract_events_from_csv(self, df: pd.DataFrame) -> List[Dict[str, Any]]:
        """
        Extract events from CSV data
        """
        events = []
        
        # Look for event-related columns
        event_columns = []
        time_columns = []
        
        for col in df.columns:
            col_lower = col.lower()
            if 'event' in col_lower:
                event_columns.append(col)
            elif any(time_word in col_lower for time_word in ['time', 'start', 'end', 'date']):
                time_columns.append(col)
        
        # Process each row for events
        for _, row in df.iterrows():
            event_name = None
            start_time = None
            end_time = None
            
            # Get event name
            for col in event_columns:
                if pd.notna(row[col]):
                    event_name = str(row[col])
                    break
            
            # Get times
            for col in time_columns:
                if pd.notna(row[col]):
                    if 'start' in col.lower():
                        start_time = str(row[col])
                    elif 'end' in col.lower():
                        end_time = str(row[col])
                    elif start_time is None:
                        start_time = str(row[col])
                    elif end_time is None:
                        end_time = str(row[col])
            
            if event_name and (start_time or end_time):
                events.append({
                    "name": event_name,
                    "start_time": start_time or "00:00",
                    "end_time": end_time or "00:00",
                    "event_type": "other",
                    "confidence": 0.9
                })
        
        return events
